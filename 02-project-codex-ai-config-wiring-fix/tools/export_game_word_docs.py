from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"C:\02-project")
EXPORTS = [
    (ROOT / "docs" / "system-hard-rules-integrated.md", ROOT / "game word" / "系统硬规则总稿.docx"),
    (ROOT / "docs" / "custom-consort-ai-interface.md", ROOT / "game word" / "自定义剧情妃与AI接口规则.docx"),
    (ROOT / "docs" / "character-story-nodes-and-relationship-ai.md", ROOT / "game word" / "角色剧情节点与关系AI接口.docx"),
    (ROOT / "docs" / "fixed-romance-npc-profiles.md", ROOT / "game word" / "固定可攻略NPC人设.docx"),
    (ROOT / "docs" / "palace-strife-architecture.md", ROOT / "game word" / "宫斗事务硬规则.docx"),
    (ROOT / "docs" / "emperor-behavior-architecture.md", ROOT / "game word" / "皇帝行为与心情硬规则.docx"),
    (ROOT / "docs" / "imperial-heir-architecture.md", ROOT / "game word" / "皇嗣管理与生育后续硬规则.docx"),
    (ROOT / "docs" / "rank-governance-architecture.md", ROOT / "game word" / "晋升降位冷宫与协理六宫硬规则.docx"),
    (ROOT / "docs" / "economy-governance-architecture.md", ROOT / "game word" / "经济与案件银两干预硬规则.docx"),
    (ROOT / "docs" / "nightly-pregnancy-architecture.md", ROOT / "game word" / "侍寝与怀孕硬规则.docx"),
    (ROOT / "reports" / "game-architecture.txt", ROOT / "game word" / "游戏架构目录.docx"),
]


def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        style.font.size = Pt(10)
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    header = rows[0]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, cell in enumerate(header):
        hdr_cells[idx].text = cell
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, cell in enumerate(row):
            cells[idx].text = cell


def parse_table_line(line: str) -> list[str]:
    parts = [part.strip() for part in line.strip().strip("|").split("|")]
    return parts


def render_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    in_code = False
    table_rows: list[list[str]] = []
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table(doc, table_rows)
            table_rows = []
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(style="CodeBlock")
            p.add_run(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            parsed = parse_table_line(line)
            if set("".join(parsed)) <= {"-", ":"}:
                continue
            table_rows.append(parsed)
            continue
        flush_table(doc, table_rows)
        table_rows = []
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("#"):
            level = min(line.count("#"), 4)
            doc.add_heading(line[level:].strip(), level=level)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line[:2].isdigit() and line[1:3] == ". ":
            doc.add_paragraph(line[3:].strip(), style="List Number")
            continue
        doc.add_paragraph(line)
    flush_table(doc, table_rows)


def render_plain(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.endswith("：") or line.endswith(":"):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)


def export_one(src: Path, dst: Path) -> None:
    doc = Document()
    ensure_styles(doc)
    text = src.read_text(encoding="utf-8")
    if src.suffix.lower() == ".md":
        render_markdown(doc, text)
    else:
        render_plain(doc, text)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)


def main() -> None:
    for src, dst in EXPORTS:
        try:
            export_one(src, dst)
            print(f"exported {dst}")
        except PermissionError:
            fallback = dst.with_name(f"{dst.stem}-更新版{dst.suffix}")
            export_one(src, fallback)
            print(f"exported {fallback}")


if __name__ == "__main__":
    main()
