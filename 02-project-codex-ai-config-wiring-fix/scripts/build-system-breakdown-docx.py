from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\02-project-codex-ai-config-wiring-fix\02-project-codex-ai-config-wiring-fix")
SRC = ROOT / "docs" / "game-system-breakdown.md"
OUT = ROOT / "docs" / "exports" / "game-system-breakdown.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Heading 1", 18, "17324D"),
        ("Heading 2", 14, "1D4E89"),
        ("Heading 3", 12, "315F72"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        code.font.size = Pt(8)
        code.paragraph_format.left_indent = Cm(0.3)
        code.paragraph_format.right_indent = Cm(0.1)
        code.paragraph_format.space_after = Pt(0)


def add_cover(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("后宫养成游戏系统拆解文档")
    set_run_font(run, 24, True, "17324D")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("系统策划版 | 架构、循环、数值与 AI 边界")
    set_run_font(run, 13, False, "52616B")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("来源：game word 规则文档、reports/game-architecture.txt、docs/*architecture.md、src/game、src/config、server/src")
    set_run_font(run, 9, False, "6B7280")

    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= {"-", ":"} for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = table.rows[row_index].cells if row_index == 0 else table.add_row().cells
        for col_index in range(col_count):
            text = row[col_index] if col_index < len(row) else ""
            cell = cells[col_index]
            cell.text = text.replace("<br>", "\n")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "D8E8F3")
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, 9, True, "17324D")
            else:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, 8.5, False, "111827")
                if row_index % 2 == 0:
                    set_cell_shading(cell, "F7FAFC")
    doc.add_paragraph()


def strip_inline_code(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def add_mermaid_block(doc: Document, code: list[str], idx: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    run = p.add_run(f"图表 {idx}：Mermaid 图源码")
    set_run_font(run, 9, True, "315F72")

    for line in code:
        p = doc.add_paragraph(style="CodeBlock")
        run = p.add_run(line)
        set_run_font(run, 7.5, False, "374151", "Consolas")


def build_docx() -> None:
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    mermaid_idx = 1
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip().lstrip("`").strip()
                code_lines = []
            else:
                if code_lang == "mermaid":
                    add_mermaid_block(doc, code_lines, mermaid_idx)
                    mermaid_idx += 1
                else:
                    for code_line in code_lines:
                        p = doc.add_paragraph(style="CodeBlock")
                        run = p.add_run(code_line)
                        set_run_font(run, 8, False, "374151", "Consolas")
                in_code = False
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            rows, next_i = parse_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(strip_inline_code(line.lstrip("> ").strip()))
            set_run_font(run, 9, False, "52616B")
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(strip_inline_code(line[2:].strip()), 1)
        elif line.startswith("## "):
            doc.add_heading(strip_inline_code(line[3:].strip()), 2)
        elif line.startswith("### "):
            doc.add_heading(strip_inline_code(line[4:].strip()), 3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(strip_inline_code(line[2:].strip()))
        else:
            p = doc.add_paragraph()
            p.add_run(strip_inline_code(line))
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
